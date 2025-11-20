"""
Enhanced converter with better image handling and professional formatting
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from PIL import Image
import os

def convert_with_images():
    """Convert with proper image handling"""
    
    # Read the markdown file
    with open("Simplified_Project_Proposal.md", 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set up styles and formatting
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1) 
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
    
    # Modify styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    
    # Process content
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Main title
        if line.startswith('# '):
            title = line[2:].strip()
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Format title
            for run in title_para.runs:
                run.font.size = Pt(18)
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0, 70, 131)  # Professional blue
        
        # Department info (bold lines after title)
        elif line.startswith('**Department') or line.startswith('**Supervisor') or line.startswith('**Team') or line.startswith('**Project Title'):
            text = line[2:-2] if line.startswith('**') and line.endswith('**') else line
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.bold = True
        
        # Section headings
        elif line.startswith('## '):
            heading = line[3:].strip()
            h = doc.add_heading(heading, level=2)
            for run in h.runs:
                run.font.size = Pt(14)
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0, 70, 131)
        
        # Subsection headings  
        elif line.startswith('### '):
            heading = line[4:].strip()
            h = doc.add_heading(heading, level=3)
            for run in h.runs:
                run.font.size = Pt(12)
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0, 70, 131)
        
        # Sub-subsection headings
        elif line.startswith('#### '):
            heading = line[5:].strip()
            h = doc.add_heading(heading, level=4)
            for run in h.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
        
        # Images with better error handling
        elif line.startswith('!['):
            match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if match:
                alt_text = match.group(1)
                image_path = match.group(2)
                
                # Check for caption on next line
                caption = ""
                if i + 1 < len(lines) and lines[i + 1].strip().startswith('*'):
                    caption = lines[i + 1].strip()[1:-1]
                    i += 1
                
                # Try to add image with fallback
                image_added = False
                try:
                    if Path(image_path).exists():
                        # Check if image is valid
                        try:
                            with Image.open(image_path) as img:
                                width, height = img.size
                                # Calculate appropriate size (max 6 inches wide)
                                max_width = 6.0
                                if width > height:
                                    img_width = min(max_width, width / 100.0)  # Rough conversion
                                else:
                                    img_width = min(max_width * 0.8, width / 100.0)
                                
                                para = doc.add_paragraph()
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = para.add_run()
                                run.add_picture(image_path, width=Inches(img_width))
                                image_added = True
                                
                        except Exception as img_error:
                            print(f"Image format issue with {image_path}: {img_error}")
                except Exception as e:
                    print(f"Could not process image {image_path}: {e}")
                
                # Add caption or placeholder
                if caption or alt_text:
                    cap_para = doc.add_paragraph()
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if image_added:
                        cap_text = caption or f"Figure: {alt_text}"
                    else:
                        cap_text = f"[Image Placeholder: {caption or alt_text}]"
                    
                    cap_run = cap_para.add_run(cap_text)
                    cap_run.italic = True
                    cap_run.font.size = Pt(10)
                    cap_run.font.name = 'Calibri'
                    cap_run.font.color.rgb = RGBColor(80, 80, 80)
        
        # Tables with better formatting
        elif '|' in line and line.count('|') >= 2:
            table_rows = []
            while i < len(lines) and '|' in lines[i]:
                row = lines[i].strip()
                if not ('---' in row):
                    table_rows.append(row)
                i += 1
            i -= 1
            
            if len(table_rows) > 0:
                first_row_cells = [cell.strip() for cell in table_rows[0].split('|') if cell.strip()]
                cols = len(first_row_cells)
                
                if cols > 0:
                    table = doc.add_table(rows=len(table_rows), cols=cols)
                    table.style = 'Table Grid'
                    
                    for row_idx, row_text in enumerate(table_rows):
                        cells = [cell.strip() for cell in row_text.split('|') if cell.strip()]
                        for col_idx, cell_text in enumerate(cells):
                            if col_idx < cols and row_idx < len(table.rows):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_text
                                # Style cells
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = 'Calibri'
                                        run.font.size = Pt(10)
                                        if row_idx == 0:  # Header row
                                            run.bold = True
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                # Add some spacing
                doc.add_paragraph()
                code_para = doc.add_paragraph()
                code_text = '\n'.join(code_lines)
                code_run = code_para.add_run(code_text)
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(9)
                code_run.font.color.rgb = RGBColor(60, 60, 60)
                doc.add_paragraph()  # Add spacing after
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            para = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
        
        elif re.match(r'^\d+\. ', line):
            list_text = re.sub(r'^\d+\. ', '', line).strip()
            para = doc.add_paragraph(list_text, style='List Number')
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
        
        # Horizontal rules
        elif line.startswith('---'):
            # Add page break for major sections
            doc.add_page_break()
        
        # Regular paragraphs with better formatting
        else:
            if line:
                para = doc.add_paragraph()
                
                # Handle bold text
                if line.startswith('**') and line.endswith('**') and len(line) > 4:
                    text = line[2:-2]
                    run = para.add_run(text)
                    run.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                else:
                    # Basic inline formatting
                    parts = re.split(r'(\*\*[^*]+\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = para.add_run(part[2:-2])
                            run.bold = True
                        elif part:
                            run = para.add_run(part)
                        
                        if 'run' in locals():
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
        
        i += 1
    
    # Save with professional name
    output_file = "Enhanced_Project_Proposal_Final.docx"
    doc.save(output_file)
    print(f"Successfully created: {output_file}")
    print(f"Document includes proper formatting, tables, and image placeholders")
    return output_file

if __name__ == "__main__":
    # Install PIL if needed
    try:
        from PIL import Image
    except ImportError:
        print("Installing Pillow for better image handling...")
        os.system("pip install Pillow")
        from PIL import Image
    
    convert_with_images()