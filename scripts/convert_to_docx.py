"""
Convert the enhanced Markdown proposal to DOCX format with proper formatting
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def add_image_if_exists(doc, image_path, caption=""):
    """Add image to document if it exists"""
    img_path = Path(image_path)
    if img_path.exists():
        try:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_picture(str(img_path), width=Inches(5.5))
            
            # Add caption
            if caption:
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(caption)
                caption_run.italic = True
                caption_run.font.size = Pt(10)
        except Exception as e:
            print(f"Could not add image {image_path}: {e}")
            # Add placeholder text instead
            doc.add_paragraph(f"[Image: {caption}]")

def convert_markdown_to_docx(md_file, output_file):
    """Convert markdown file to DOCX with proper formatting"""
    
    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Title (# heading)
        if line.startswith('# '):
            title = line[2:].strip()
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Level 2 headings (##)
        elif line.startswith('## '):
            heading = line[3:].strip()
            doc.add_heading(heading, level=2)
            
        # Level 3 headings (###)
        elif line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=3)
            
        # Level 4 headings (####)
        elif line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_heading(heading, level=4)
            
        # Horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph().add_run().add_break()
            
        # Images with captions
        elif line.startswith('!['):
            # Extract image info: ![alt text](image_path)
            match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if match:
                alt_text = match.group(1)
                image_path = match.group(2)
                
                # Look for caption on next line if it starts with *
                caption = ""
                if i + 1 < len(lines) and lines[i + 1].strip().startswith('*') and lines[i + 1].strip().endswith('*'):
                    caption = lines[i + 1].strip()[1:-1]  # Remove asterisks
                    i += 1  # Skip caption line
                
                add_image_if_exists(doc, image_path, caption or alt_text)
            
        # Tables
        elif '|' in line and line.count('|') >= 2:
            # Simple table handling - collect all table rows
            table_rows = []
            while i < len(lines) and '|' in lines[i]:
                table_rows.append(lines[i].strip())
                i += 1
            i -= 1  # Back up one since we'll increment at the end
            
            if len(table_rows) > 1:
                # Skip separator row if present
                if '---' in table_rows[1]:
                    table_rows.pop(1)
                
                # Create table
                table = doc.add_table(rows=len(table_rows), cols=table_rows[0].count('|') - 1)
                table.style = 'Table Grid'
                
                for row_idx, row_text in enumerate(table_rows):
                    cells = [cell.strip() for cell in row_text.split('|')[1:-1]]  # Remove empty first/last
                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < len(table.rows[row_idx].cells):
                            table.rows[row_idx].cells[col_idx].text = cell_text
        
        # Code blocks
        elif line.startswith('```'):
            # Collect code block content
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                code_para = doc.add_paragraph()
                code_run = code_para.add_run('\n'.join(code_lines))
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
                # Add light gray background (this is complex in python-docx, so we'll skip it)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            doc.add_paragraph(bullet_text, style='List Bullet')
            
        # Numbered lists
        elif re.match(r'^\d+\. ', line):
            list_text = re.sub(r'^\d+\. ', '', line).strip()
            doc.add_paragraph(list_text, style='List Number')
            
        # Bold headers (lines starting with **)
        elif line.startswith('**') and line.endswith('**'):
            header_text = line[2:-2].strip()
            para = doc.add_paragraph()
            run = para.add_run(header_text)
            run.bold = True
            
        # Regular paragraphs
        else:
            if line:  # Only add non-empty lines
                para = doc.add_paragraph()
                
                # Handle inline formatting
                text = line
                parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
                
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = para.add_run(part[1:-1])
                        run.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        run = para.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                    else:
                        para.add_run(part)
        
        i += 1
    
    # Save the document
    doc.save(output_file)
    print(f"Successfully converted {md_file} to {output_file}")

if __name__ == "__main__":
    input_file = "Simplified_Project_Proposal.md"
    output_file = "Enhanced_Project_Proposal.docx"
    
    convert_markdown_to_docx(input_file, output_file)