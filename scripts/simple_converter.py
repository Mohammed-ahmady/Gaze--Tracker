"""
Simple and reliable Markdown to DOCX converter for the project proposal
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx():
    """Convert the enhanced proposal to DOCX"""
    
    # Read the markdown file
    with open("Simplified_Project_Proposal.md", 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Main title
        if line.startswith('# '):
            title = line[2:].strip()
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_para.runs:
                run.font.size = Pt(16)
                run.font.name = 'Calibri'
        
        # Section headings
        elif line.startswith('## '):
            heading = line[3:].strip()
            h = doc.add_heading(heading, level=2)
            for run in h.runs:
                run.font.size = Pt(14)
                run.font.name = 'Calibri'
        
        # Subsection headings
        elif line.startswith('### '):
            heading = line[4:].strip()
            h = doc.add_heading(heading, level=3)
            for run in h.runs:
                run.font.size = Pt(12)
                run.font.name = 'Calibri'
        
        # Sub-subsection headings
        elif line.startswith('#### '):
            heading = line[5:].strip()
            h = doc.add_heading(heading, level=4)
            for run in h.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
        
        # Images
        elif line.startswith('!['):
            match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if match:
                alt_text = match.group(1)
                image_path = match.group(2)
                
                # Check for caption on next line
                caption = ""
                if i + 1 < len(lines) and lines[i + 1].strip().startswith('*'):
                    caption = lines[i + 1].strip()[1:-1]  # Remove asterisks
                    i += 1  # Skip caption line
                
                # Try to add image
                try:
                    if Path(image_path).exists():
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        run.add_picture(image_path, width=Inches(5.5))
                        
                        # Add caption
                        if caption:
                            cap_para = doc.add_paragraph()
                            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap_run = cap_para.add_run(caption)
                            cap_run.italic = True
                            cap_run.font.size = Pt(10)
                            cap_run.font.name = 'Calibri'
                    else:
                        # Add placeholder
                        para = doc.add_paragraph(f"[Figure: {alt_text}]")
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.italic = True
                            run.font.name = 'Calibri'
                except Exception as e:
                    print(f"Could not add image {image_path}: {e}")
                    para = doc.add_paragraph(f"[Figure: {alt_text}]")
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tables
        elif '|' in line and line.count('|') >= 2:
            # Collect table rows
            table_rows = []
            start_i = i
            while i < len(lines) and '|' in lines[i]:
                row = lines[i].strip()
                if not ('---' in row):  # Skip separator rows
                    table_rows.append(row)
                i += 1
            i -= 1  # Back up one
            
            if len(table_rows) > 0:
                # Parse first row to get column count
                first_row_cells = [cell.strip() for cell in table_rows[0].split('|') if cell.strip()]
                cols = len(first_row_cells)
                
                if cols > 0:
                    table = doc.add_table(rows=len(table_rows), cols=cols)
                    table.style = 'Table Grid'
                    
                    for row_idx, row_text in enumerate(table_rows):
                        cells = [cell.strip() for cell in row_text.split('|') if cell.strip()]
                        for col_idx, cell_text in enumerate(cells):
                            if col_idx < cols and row_idx < len(table.rows):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_text
                                # Make header row bold
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                code_para = doc.add_paragraph()
                code_text = '\n'.join(code_lines)
                code_run = code_para.add_run(code_text)
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            para = doc.add_paragraph(bullet_text, style='List Bullet')
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
        
        # Numbered lists
        elif re.match(r'^\d+\. ', line):
            list_text = re.sub(r'^\d+\. ', '', line).strip()
            para = doc.add_paragraph(list_text, style='List Number')
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
        
        # Horizontal rules (skip)
        elif line.startswith('---'):
            doc.add_paragraph()  # Just add space
        
        # Regular paragraphs
        else:
            if line:
                para = doc.add_paragraph()
                
                # Simple bold formatting
                if line.startswith('**') and line.endswith('**') and len(line) > 4:
                    text = line[2:-2]
                    run = para.add_run(text)
                    run.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                else:
                    # Handle inline formatting (basic)
                    text = line
                    # Split by **bold** patterns
                    parts = re.split(r'(\*\*[^*]+\*\*)', text)
                    
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = para.add_run(part[2:-2])
                            run.bold = True
                        elif part:
                            run = para.add_run(part)
                        
                        if 'run' in locals():
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
        
        i += 1
    
    # Save the document
    output_file = "Enhanced_Project_Proposal.docx"
    doc.save(output_file)
    print(f"Successfully converted to {output_file}")
    return output_file

if __name__ == "__main__":
    convert_md_to_docx()